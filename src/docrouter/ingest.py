"""Turning files on disk into `Document` objects.

Two registries carry the extensibility here. `ExtractorRegistry` maps file
suffixes to extractors, and `OcrRegistry` maps names to OCR backends. Adding
a format or a transcription engine means writing a class and registering it;
no dispatch chain in this module gets edited, and no caller changes. That is
the open/closed boundary, and `test_ingest.py` exercises it by registering a
throwaway extractor at runtime.

The mixed scanned/digital problem is handled inside `PdfExtractor`, per page:
government PDFs are routinely a digital cover sheet stapled to scanned
attachments, so deciding once per file would OCR far too much or far too
little.
"""

from __future__ import annotations

import base64
import io
import os
from collections.abc import Callable, Sequence
from dataclasses import dataclass
from pathlib import Path

from . import normalize
from .models import Document, Source
from .protocols import OcrBackend, TextExtractor

# Below this many characters, a page's text layer is treated as absent.
MIN_CHARS_PER_PAGE = 40
# Rasterization DPI for OCR. 200 is the floor for reliable Arabic diacritics.
OCR_DPI = 200

IMAGE_SUFFIXES = frozenset({".png", ".jpg", ".jpeg", ".tif", ".tiff"})


class OCRUnavailable(RuntimeError):
    """Raised when a requested OCR backend cannot run in this environment."""


class UnsupportedDocument(ValueError):
    """Raised for a file type no registered extractor claims."""


@dataclass(frozen=True)
class Extraction:
    """Text pulled from one file, plus how it was obtained."""

    text: str
    source: Source
    page_count: int


# --------------------------------------------------------------------------
# OCR backends
# --------------------------------------------------------------------------

OCR_INSTRUCTION = (
    "انسخ كل النص الظاهر في هذه الصفحات حرفياً كما هو، بالترتيب.\n"
    "- لا تترجم ولا تلخّص ولا تصحح الأخطاء الإملائية.\n"
    "- حافظ على الأسطر والعناوين وأرقام الجداول.\n"
    "- إن كانت صفحة غير مقروءة تماماً فاكتب [صفحة غير مقروءة].\n"
    "أخرج النص فقط دون أي تعليق."
)


class ClaudeVisionOcr:
    """Default backend: no external binaries, strong on Arabic ligatures.

    The client is injected rather than constructed here so tests and offline
    callers can supply their own, and so no API object is created until a
    document actually needs transcribing.
    """

    def __init__(self, client=None, model: str | None = None, max_tokens: int = 16000):
        self._client = client
        self._model = model
        self.max_tokens = max_tokens

    @property
    def name(self) -> str:
        return "claude"

    def _ensure_client(self):
        if self._client is None:
            import anthropic

            self._client = anthropic.Anthropic()
        return self._client

    def _resolve_model(self) -> str:
        from . import DEFAULT_MODEL

        return self._model or os.environ.get("DOCROUTER_MODEL", DEFAULT_MODEL)

    def transcribe(self, images: Sequence[bytes]) -> str:
        if not images:
            return ""
        content: list[dict] = [
            {
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": "image/png",
                    "data": base64.standard_b64encode(image).decode("ascii"),
                },
            }
            for image in images
        ]
        content.append({"type": "text", "text": OCR_INSTRUCTION})

        response = self._ensure_client().messages.create(
            model=self._resolve_model(),
            max_tokens=self.max_tokens,
            messages=[{"role": "user", "content": content}],
        )
        return "".join(b.text for b in response.content if b.type == "text")


class TesseractOcr:
    """Local alternative for material that must not leave the network."""

    def __init__(self, lang: str = "ara+eng"):
        self.lang = lang

    @property
    def name(self) -> str:
        return "tesseract"

    def transcribe(self, images: Sequence[bytes]) -> str:
        if not images:
            return ""
        try:
            import pytesseract
            from PIL import Image
        except ImportError as exc:  # pragma: no cover - optional extra
            raise OCRUnavailable(
                "tesseract backend needs `pip install 'docrouter[tesseract]'` "
                "plus the Tesseract binary with the `ara` language pack"
            ) from exc

        return "\n\n".join(
            pytesseract.image_to_string(Image.open(io.BytesIO(image)), lang=self.lang)
            for image in images
        )


class OcrRegistry:
    """Name -> OCR backend factory."""

    def __init__(self) -> None:
        self._factories: dict[str, Callable[..., OcrBackend]] = {}

    def register(self, name: str, factory: Callable[..., OcrBackend]) -> None:
        self._factories[name] = factory

    def create(self, name: str, **kwargs) -> OcrBackend:
        try:
            factory = self._factories[name]
        except KeyError:
            raise ValueError(
                f"unknown OCR backend {name!r}; registered: {sorted(self._factories)}"
            ) from None
        return factory(**kwargs)

    @property
    def names(self) -> frozenset[str]:
        return frozenset(self._factories)


OCR_REGISTRY = OcrRegistry()
OCR_REGISTRY.register("claude", ClaudeVisionOcr)
OCR_REGISTRY.register(
    "tesseract", lambda **kw: TesseractOcr(**{k: v for k, v in kw.items() if k == "lang"})
)


def resolve_ocr(backend: OcrBackend | str | None, client=None) -> OcrBackend | None:
    """Accept a backend instance, a registered name, or None."""
    if backend is None or isinstance(backend, (ClaudeVisionOcr, TesseractOcr)):
        return backend
    if isinstance(backend, str):
        if backend == "claude":
            return OCR_REGISTRY.create("claude", client=client)
        return OCR_REGISTRY.create(backend)
    return backend  # already satisfies the protocol


# --------------------------------------------------------------------------
# Extractors
# --------------------------------------------------------------------------


class PlainTextExtractor:
    @property
    def suffixes(self) -> frozenset[str]:
        return frozenset({".txt", ".md"})

    def extract(self, path: Path, ocr: OcrBackend | None = None) -> Extraction:
        return Extraction(path.read_text(encoding="utf-8"), "digital", 1)


class DocxExtractor:
    @property
    def suffixes(self) -> frozenset[str]:
        return frozenset({".docx"})

    def extract(self, path: Path, ocr: OcrBackend | None = None) -> Extraction:
        import docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return Extraction("\n".join(parts), "digital", 1)


class ImageExtractor:
    @property
    def suffixes(self) -> frozenset[str]:
        return IMAGE_SUFFIXES

    def extract(self, path: Path, ocr: OcrBackend | None = None) -> Extraction:
        if ocr is None:
            raise OCRUnavailable(f"{path.name} is an image and needs an OCR backend")
        return Extraction(ocr.transcribe([path.read_bytes()]), "ocr", 1)


def arabic_text_layer_is_usable(text: str, min_chars: int = MIN_CHARS_PER_PAGE) -> bool:
    """Default policy: enough characters, and actually Arabic.

    The second half matters as much as the first. A page whose text layer is
    Latin-only is the signature of a scanner that exported garbage, and it
    should go to OCR even though it is not empty.
    """
    return len(text.strip()) >= min_chars and normalize.looks_arabic(text)


class PdfExtractor:
    """Reads the embedded text layer, OCR-ing only the pages that lack one.

    `is_usable` is injectable because "usable" is a deployment policy, not a
    fact about PDFs: an Arabic archive wants the default, a bilingual or
    Latin-script one does not.
    """

    def __init__(
        self,
        min_chars: int = MIN_CHARS_PER_PAGE,
        dpi: int = OCR_DPI,
        is_usable: Callable[[str], bool] | None = None,
    ):
        self.min_chars = min_chars
        self.dpi = dpi
        self._is_usable = is_usable or (
            lambda text: arabic_text_layer_is_usable(text, self.min_chars)
        )

    @property
    def suffixes(self) -> frozenset[str]:
        return frozenset({".pdf"})

    def _needs_ocr(self, text: str) -> bool:
        return not self._is_usable(text)

    def extract(self, path: Path, ocr: OcrBackend | None = None) -> Extraction:
        import fitz  # PyMuPDF

        document = fitz.open(path)
        try:
            page_texts: list[str] = []
            scanned: list[int] = []

            for index, page in enumerate(document):
                text = page.get_text("text") or ""
                if self._needs_ocr(text):
                    scanned.append(index)
                    page_texts.append("")
                else:
                    page_texts.append(text)

            if scanned:
                if ocr is None:
                    raise OCRUnavailable(
                        f"{path.name} has {len(scanned)} page(s) without a usable "
                        "text layer and no OCR backend was supplied"
                    )
                zoom = self.dpi / 72
                images = [
                    document[i].get_pixmap(matrix=fitz.Matrix(zoom, zoom)).tobytes("png")
                    for i in scanned
                ]
                # The backend transcribes the batch as one stream; attaching it
                # to the first scanned page keeps document order close enough
                # for classification, which reads the whole document anyway.
                page_texts[scanned[0]] = ocr.transcribe(images)

            return Extraction(
                "\n\n".join(t for t in page_texts if t),
                "ocr" if scanned else "digital",
                len(page_texts),
            )
        finally:
            document.close()


class ExtractorRegistry:
    """Suffix -> extractor. Later registrations win, so defaults are replaceable."""

    def __init__(self, extractors: Sequence[TextExtractor] = ()) -> None:
        self._by_suffix: dict[str, TextExtractor] = {}
        for extractor in extractors:
            self.register(extractor)

    def register(self, extractor: TextExtractor) -> None:
        for suffix in extractor.suffixes:
            self._by_suffix[suffix.lower()] = extractor

    def get(self, suffix: str) -> TextExtractor:
        try:
            return self._by_suffix[suffix.lower()]
        except KeyError:
            raise UnsupportedDocument(
                f"unsupported file type {suffix!r}; supported: {sorted(self.suffixes)}"
            ) from None

    def supports(self, suffix: str) -> bool:
        return suffix.lower() in self._by_suffix

    @property
    def suffixes(self) -> frozenset[str]:
        return frozenset(self._by_suffix)


def default_registry() -> ExtractorRegistry:
    return ExtractorRegistry(
        [PdfExtractor(), DocxExtractor(), ImageExtractor(), PlainTextExtractor()]
    )


DEFAULT_REGISTRY = default_registry()


def supported_suffixes(registry: ExtractorRegistry | None = None) -> frozenset[str]:
    return (registry or DEFAULT_REGISTRY).suffixes


# Kept for callers that predate the registry.
SUPPORTED_SUFFIXES = DEFAULT_REGISTRY.suffixes


# --------------------------------------------------------------------------
# Public entry points
# --------------------------------------------------------------------------


def run_ocr(images: Sequence[bytes], backend: OcrBackend | str = "claude", client=None) -> str:
    """Transcribe page images with a named or supplied backend."""
    resolved = resolve_ocr(backend, client=client)
    if resolved is None:
        raise OCRUnavailable("no OCR backend supplied")
    return resolved.transcribe(images)


def load_document(
    path: str | Path,
    ocr_backend: OcrBackend | str | None = "claude",
    client=None,
    registry: ExtractorRegistry | None = None,
) -> Document:
    """Ingest one file, OCR-ing only what needs it."""
    path = Path(path)
    registry = registry or DEFAULT_REGISTRY
    extractor = registry.get(path.suffix)
    ocr = resolve_ocr(ocr_backend, client=client)

    result = extractor.extract(path, ocr)
    return Document(
        doc_id=path.stem,
        text=normalize.light(result.text),
        source=result.source,
        path=str(path),
        page_count=result.page_count,
    )


def load_directory(
    directory: str | Path,
    ocr_backend: OcrBackend | str | None = "claude",
    client=None,
    registry: ExtractorRegistry | None = None,
) -> list[Document]:
    """Ingest every supported file under a directory, recursively."""
    directory = Path(directory)
    registry = registry or DEFAULT_REGISTRY
    ocr = resolve_ocr(ocr_backend, client=client)

    files = sorted(p for p in directory.rglob("*") if registry.supports(p.suffix))
    return [load_document(p, ocr_backend=ocr, registry=registry) for p in files]
